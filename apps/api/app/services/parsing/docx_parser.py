"""DOCX parser: extracts paragraph text using python-docx. No images or OLE objects."""
from app.services.parsing.base import MAX_TEXT_BYTES, CURRENT_PARSER_VERSION, ParserResult

PARSER_NAME = "docx"


def parse_docx(content: bytes) -> ParserResult:
    try:
        import io
        from docx import Document  # type: ignore[import-untyped]
    except ImportError:
        return ParserResult(
            text="", parser_name=PARSER_NAME, parser_version=CURRENT_PARSER_VERSION,
            error="python-docx not installed — DOCX parsing is not available",
        )

    try:
        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n".join(paragraphs)

        if len(text.encode("utf-8")) > MAX_TEXT_BYTES:
            text = text.encode("utf-8")[:MAX_TEXT_BYTES].decode("utf-8", errors="ignore")

        return ParserResult(
            text=text,
            parser_name=PARSER_NAME,
            parser_version=CURRENT_PARSER_VERSION,
            metadata_json={"paragraph_count": len(doc.paragraphs)},
        )
    except Exception as exc:
        return ParserResult(
            text="", parser_name=PARSER_NAME, parser_version=CURRENT_PARSER_VERSION,
            error=f"DOCX parse error: {str(exc)[:500]}",
        )
